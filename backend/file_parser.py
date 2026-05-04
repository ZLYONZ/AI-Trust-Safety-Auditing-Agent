"""
file_parser.py
--------------
Extract plain text from uploaded audit documents, URLs, and GitHub repos.
 
Supported formats
  .pdf   — pdfplumber (text-layer PDFs) with pdfminer fallback
  .docx  — python-docx
  .txt / .md / .csv / .json / .xml  — read as UTF-8
  .xlsx  — openpyxl
  URL    — fetches HTML/PDF/text from public URLs
  GitHub — fetches README, governance docs, and markdown from public repos
 
Install dependencies:
  pip install pdfplumber python-docx openpyxl
"""
 
import io
import json
import re
import ssl
import urllib.request
import urllib.error
from pathlib import Path
 
# Fix SSL certificate verification on macOS
# Equivalent to running Install Certificates.command
try:
    import certifi
    _SSL_CONTEXT = ssl.create_default_context(cafile=certifi.where())
except ImportError:
    _SSL_CONTEXT = ssl.create_default_context()
 
 
def _open_url(req: urllib.request.Request, timeout: int = 15):
    """Open a URL with SSL cert fix applied."""
    return urllib.request.urlopen(req, timeout=timeout, context=_SSL_CONTEXT)
 
 
# ── File extraction ──────────────────────────────────────────────────────────

def extract_text(file_bytes: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(file_bytes, filename)
    elif suffix in (".docx", ".doc"):
        return _extract_docx(file_bytes, filename)
    elif suffix in (".txt", ".md"):
        return file_bytes.decode("utf-8", errors="replace")
    elif suffix == ".csv":
        return _extract_csv(file_bytes)
    elif suffix == ".json":
        return _extract_json(file_bytes)
    elif suffix in (".xml", ".html", ".htm"):
        return _strip_xml_tags(file_bytes.decode("utf-8", errors="replace"))
    elif suffix in (".xlsx", ".xls"):
        return _extract_xlsx(file_bytes, filename)
    else:
        raise ValueError(
            f"Unsupported file type '{suffix}' for '{filename}'. "
            "Accepted: pdf, docx, txt, md, csv, json, xml, xlsx."
        )


def extract_text_from_files(files: list[tuple[bytes, str]]) -> str:
    parts: list[str] = []
    for file_bytes, filename in files:
        try:
            text = extract_text(file_bytes, filename)
            parts.append(f"=== Document: {filename} ===\n{text.strip()}")
        except ValueError as exc:
            parts.append(f"=== Document: {filename} (SKIPPED — {exc}) ===")
    return "\n\n".join(parts)


# ── URL fetching ─────────────────────────────────────────────────────────────

def fetch_url_text(url: str) -> str:
    """
    Fetch text from a public URL. Tries multiple strategies:
    1. Direct fetch with browser-like headers
    2. Jina Reader API (r.jina.ai) for JS-heavy pages that block bots
    3. PDF extraction if content-type is PDF
    """
    # Strategy 1: direct fetch with realistic browser headers
    browser_headers = {
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.5",
    }

    raw_bytes    = b""
    content_type = ""

    try:
        req = urllib.request.Request(url, headers=browser_headers)
        with urllib.request.urlopen(req, timeout=15) as resp:
            content_type = resp.headers.get("Content-Type", "")
            raw_bytes    = resp.read()
    except urllib.error.HTTPError as e:
        if e.code in (403, 429, 503):
            # Strategy 2: try Jina Reader which renders JS pages
            jina_url = f"https://r.jina.ai/{url}"
            try:
                req2 = urllib.request.Request(
                    jina_url,
                    headers={"Accept": "text/plain", "User-Agent": "TrustGuard/1.0"},
                )
                with urllib.request.urlopen(req2, timeout=20) as resp2:
                    return resp2.read().decode("utf-8", errors="replace")
            except Exception:
                pass
        raise ValueError(f"Could not fetch URL '{url}': HTTP {e.code} {e.reason}")
    except urllib.error.URLError as e:
        raise ValueError(f"Could not fetch URL '{url}': {e}")

    if "pdf" in content_type or url.lower().endswith(".pdf"):
        return extract_text(raw_bytes, "fetched.pdf")
    elif "html" in content_type or b"<html" in raw_bytes[:200].lower():
        return _strip_xml_tags(raw_bytes.decode("utf-8", errors="replace"))
    else:
        return raw_bytes.decode("utf-8", errors="replace")


# ── GitHub repo fetching ──────────────────────────────────────────────────────

GITHUB_TARGET_FILES = [
    "README.md", "README.rst", "README.txt",
    "GOVERNANCE.md", "SECURITY.md", "AI_POLICY.md",
    "AI-POLICY.md", "AI_GOVERNANCE.md", "CONTRIBUTING.md",
    "CODE_OF_CONDUCT.md", "docs/governance.md", "docs/ai-policy.md",
    "docs/security.md", "docs/fairness.md", "docs/explainability.md",
    "docs/performance.md", "docs/README.md",
    "policy.md", "compliance.md", "fairness.md",
]


def _clean_markdown(text: str) -> str:
    """Strip code blocks and keep only prose from markdown files."""
    # Remove fenced code blocks
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"`[^`]+`", "", text)  # inline code
    # Remove image tags
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    # Convert links to just text
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"", text)
    # Remove HTML comments
    text = re.sub(r"<!--[\s\S]*?-->", "", text)
    # Remove horizontal rules
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    # Clean up excessive whitespace
    text = re.sub(r"{3,}", "", text)
    # Remove lines that are purely formatting (badges, shields, etc.)
    lines = []
    for line in text.split(""):
        stripped = line.strip()
        # Skip badge lines, empty lines after stripping, pure symbol lines
        if re.match(r"^\[!\[.*\]\(https://.*badge.*\)", stripped):
            continue
        lines.append(line)
    return "".join(lines).strip()


def fetch_github_repo_text(repo: str) -> str:
    owner, name = repo.split("/", 1)
    parts: list[str] = []

    # Get default branch — use GITHUB_TOKEN if set for higher rate limits
    import os as _os
    github_token = _os.environ.get("GITHUB_TOKEN", "")
    gh_headers = {
        "User-Agent": "TrustGuard-Audit-Bot/1.0",
        "Accept": "application/vnd.github.v3+json",
    }
    if github_token:
        gh_headers["Authorization"] = f"token {github_token}"

    default_branch = "main"
    try:
        req = urllib.request.Request(
            f"https://api.github.com/repos/{owner}/{name}",
            headers=gh_headers,
        )
        with urllib.request.urlopen(req, timeout=10) as r:
            info = json.loads(r.read())
            default_branch = info.get("default_branch", "main")
            if info.get("description"):
                parts.append(
                    f"=== Repository: {owner}/{name} ===\n"
                    f"Description: {info['description']}\n"
                )
    except Exception:
        pass

    # Fetch target files
    for filepath in GITHUB_TARGET_FILES:
        raw_url = (
            f"https://raw.githubusercontent.com/{owner}/{name}"
            f"/{default_branch}/{filepath}"
        )
        try:
            req = urllib.request.Request(raw_url, headers=gh_headers)
            with urllib.request.urlopen(req, timeout=8) as r:
                file_content = r.read().decode("utf-8", errors="replace")
                cleaned = _clean_markdown(file_content)
                if cleaned.strip():
                    parts.append(f"=== File: {filepath} ===\n{cleaned.strip()}")
        except Exception:
            continue

    # Scan docs/ folder for additional .md files
    try:
        req = urllib.request.Request(
            f"https://api.github.com/repos/{owner}/{name}/contents/docs",
            headers={"User-Agent": "TrustGuard-Audit-Bot/1.0",
                     "Accept": "application/vnd.github.v3+json"},
        )
        with urllib.request.urlopen(req, timeout=8) as r:
            for item in json.loads(r.read()):
                if item.get("type") == "file" and item["name"].endswith(".md"):
                    fp = f"docs/{item['name']}"
                    if any(fp in p for p in parts):
                        continue
                    try:
                        req2 = urllib.request.Request(
                            f"https://raw.githubusercontent.com/{owner}/{name}"
                            f"/{default_branch}/{fp}",
                            headers={"User-Agent": "TrustGuard-Audit-Bot/1.0"},
                        )
                        with urllib.request.urlopen(req2, timeout=8) as r2:
                            content = r2.read().decode("utf-8", errors="replace")
                            cleaned = _clean_markdown(content)
                            if cleaned.strip():
                                parts.append(f"=== File: {fp} ===\n{cleaned.strip()}")
                    except Exception:
                        continue
    except Exception:
        pass

    if not parts:
        raise ValueError(
            f"No documentation files found in github.com/{repo}. "
            "Make sure the repository is public and contains a README.md "
            "or governance documentation."
        )

    return "\n\n".join(parts)


# ── Combined source extractor ─────────────────────────────────────────────────

def extract_text_from_sources(
    files: list[tuple[bytes, str]],
    urls: list[str] | None = None,
    github_repos: list[str] | None = None,
) -> str:
    parts: list[str] = []

    if files:
        parts.append(extract_text_from_files(files))

    for url in (urls or []):
        try:
            text = fetch_url_text(url)
            parts.append(f"=== Source URL: {url} ===\n{text.strip()}")
        except ValueError as e:
            parts.append(f"=== Source URL: {url} (FAILED — {e}) ===")

    for repo in (github_repos or []):
        try:
            text = fetch_github_repo_text(repo)
            parts.append(f"=== GitHub Repository: {repo} ===\n{text.strip()}")
        except ValueError as e:
            parts.append(f"=== GitHub Repository: {repo} (FAILED — {e}) ===")

    return "\n\n".join(parts)


# ── Format extractors ────────────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes, filename: str) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Run: pip install pdfplumber")

    text_pages: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_pages.append(page_text)

    if not text_pages:
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            return pdfminer_extract(io.BytesIO(file_bytes)).strip()
        except ImportError:
            pass
        raise ValueError(
            f"No text extracted from '{filename}'. "
            "PDF may be scanned/image-based."
        )

    return "\n\n".join(text_pages)


def _extract_docx(file_bytes: bytes, filename: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Run: pip install python-docx")

    doc        = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def _extract_csv(file_bytes: bytes) -> str:
    import csv
    reader = csv.reader(io.StringIO(file_bytes.decode("utf-8", errors="replace")))
    return "\n".join(", ".join(row) for row in reader)


def _extract_json(file_bytes: bytes) -> str:
    try:
        data = json.loads(file_bytes.decode("utf-8", errors="replace"))
        return json.dumps(data, indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        return file_bytes.decode("utf-8", errors="replace")


def _strip_xml_tags(text: str) -> str:
    """Remove HTML tags — strips nav/footer/script blocks for cleaner policy text."""
    for tag in ["script", "style", "nav", "footer", "header", "aside", "noscript"]:
        text = re.sub(
            rf"<{tag}[^>]*>.*?</{tag}>", " ", text,
            flags=re.DOTALL | re.IGNORECASE
        )
    text = re.sub(r"<[^>]+>", " ", text)
    text = (text
            .replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
            .replace("&nbsp;", " ").replace("&quot;", '"').replace("&#39;", "'"))
    text = re.sub(r"[ 	]+", " ", text)
    text = re.sub(r"{3,}", "", text)
    lines = [ln for ln in text.split("") if len(ln.strip()) > 20 or not ln.strip()]
    return "".join(lines).strip()

def _extract_xlsx(file_bytes: bytes, filename: str) -> str:
    try:
        import openpyxl
    except ImportError:
        raise ImportError("Run: pip install openpyxl")

    wb    = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                parts.append(", ".join(cells))
    return "\n".join(parts)